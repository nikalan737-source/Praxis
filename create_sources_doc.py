from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
from docx.opc.constants import RELATIONSHIP_TYPE as RT
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
title = doc.add_heading('NVDP Opportunity Analysis — Source Reference', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph('Praxis: AI-Powered Personal Health Experimentation Platform')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(11)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
subtitle.runs[0].italic = True

doc.add_paragraph('')

def add_section_header(doc, text, color=(0x1a, 0x56, 0x96)):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(*color)
    h.runs[0].font.size = Pt(14)
    return h

def add_source_block(doc, credibility_label, credibility_color, source_name, source_type, url, facts):
    """Add a formatted source block with facts"""
    # Credibility badge + source name
    p = doc.add_paragraph()
    badge = p.add_run(f'[{credibility_label}]  ')
    badge.font.bold = True
    badge.font.size = Pt(10)
    badge.font.color.rgb = credibility_color
    name_run = p.add_run(source_name)
    name_run.font.bold = True
    name_run.font.size = Pt(11)
    name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Source type + URL
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    type_run = p2.add_run(f'{source_type}  |  ')
    type_run.font.size = Pt(9)
    type_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    type_run.italic = True
    url_run = p2.add_run(url)
    url_run.font.size = Pt(9)
    url_run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)

    # Facts
    for fact in facts:
        fp = doc.add_paragraph(style='List Bullet')
        fp.paragraph_format.left_indent = Inches(0.4)
        fr = fp.add_run(fact)
        fr.font.size = Pt(10)

    doc.add_paragraph('')

# ─────────────────────────────────────────────
# SECTION 1: PEER-REVIEWED / SCHOLARLY
# ─────────────────────────────────────────────
add_section_header(doc, '✅ Peer-Reviewed & Scholarly Sources')

add_source_block(doc,
    credibility_label='PEER-REVIEWED',
    credibility_color=RGBColor(0x1e, 0x88, 0x55),
    source_name="Gen Z's Enduring Commitment to Fitness Apps — Tandfonline (2024)",
    source_type='Peer-reviewed journal article | Cogent Business & Management',
    url='https://www.tandfonline.com/doi/full/10.1080/23311975.2024.2419483',
    facts=[
        'Academic study examining Gen Z\'s sustained engagement with fitness and wellness apps',
        'Validates the behavioral trend of young adults using digital tools to manage health',
        'Access through UW Library — search Tandfonline database or DOI: 10.1080/23311975.2024.2419483',
        'Use for: backing the claim that Gen Z is the core digital health consumer segment',
    ]
)

add_source_block(doc,
    credibility_label='PEER-REVIEWED',
    credibility_color=RGBColor(0x1e, 0x88, 0x55),
    source_name='Health-Related Communication of Social Media Influencers: A Scoping Review — Tandfonline (2024)',
    source_type='Peer-reviewed journal article | Health Communication',
    url='https://www.tandfonline.com/doi/full/10.1080/10410236.2024.2397268',
    facts=[
        'Systematic review of how health influencers impact consumer health behavior',
        '82% of consumers report being likely to follow health advice from social media influencers',
        'Health influencer content receives 45% more engagement than posts from traditional health organizations',
        'Documents the credibility gap in influencer-driven health content — consumers want science-backed alternatives',
        'Access through UW Library — search DOI: 10.1080/10410236.2024.2397268',
        'Use for: the "credibility gap" argument and why Praxis fills an unmet need',
    ]
)

# ─────────────────────────────────────────────
# SECTION 2: HIGHLY CREDIBLE INDUSTRY / CONSULTING
# ─────────────────────────────────────────────
add_section_header(doc, '✅ Highly Credible — Major Consulting & Business Sources')

add_source_block(doc,
    credibility_label='STRONG',
    credibility_color=RGBColor(0x15, 0x65, 0xC0),
    source_name='The $2 Trillion Global Wellness Market — McKinsey & Company',
    source_type='Management consulting research report | McKinsey & Company',
    url='https://www.mckinsey.com/industries/consumer-packaged-goods/our-insights/future-of-wellness-trends',
    facts=[
        'Gen Z and Millennials make up 36% of U.S. adults but drive 41% of annual wellness spending',
        'Documents the generational shift in wellness spending toward younger demographics',
        'Health and wellness is consistently the second most popular category of influencer content consumers act on',
        'Use for: consumer spending data and generational market shift argument',
    ]
)

add_source_block(doc,
    credibility_label='STRONG',
    credibility_color=RGBColor(0x15, 0x65, 0xC0),
    source_name='The Trends Defining the $1.8 Trillion Global Wellness Market in 2024 — McKinsey & Company',
    source_type='Management consulting research report | McKinsey & Company',
    url='https://www.mckinsey.com/industries/consumer-packaged-goods/our-insights/the-trends-defining-the-1-point-8-trillion-dollar-global-wellness-market-in-2024',
    facts=[
        'Global wellness economy projected to reach $7 trillion in 2025',
        'Documents the "healthwashing" backlash — consumers in 2024–2025 are actively seeking science-backed, credible health guidance',
        'Personalized and preventive health is the fastest-growing wellness segment',
        'Use for: overall market size, credibility gap narrative, and "why now" argument',
    ]
)

add_source_block(doc,
    credibility_label='STRONG',
    credibility_color=RGBColor(0x15, 0x65, 0xC0),
    source_name='Is Your Gym Overflowing? Gen Z Is Driving the Trend — Fortune',
    source_type='Major business publication | Fortune Well',
    url='https://fortune.com/well/article/gym-use-increase-pandemic-gen-z/',
    facts=[
        'Gym use is nearly double pre-pandemic levels',
        'Gen Z is identified as the primary driver of the gym membership surge',
        'Cites academic and industry research on fitness participation trends',
        'Use for: establishing the cultural trend of Gen Z health engagement',
    ]
)

# ─────────────────────────────────────────────
# SECTION 3: CREDIBLE INDUSTRY / TRADE SOURCES
# ─────────────────────────────────────────────
add_section_header(doc, '✅ Credible Industry & Trade Sources (cite as "industry data")')

add_source_block(doc,
    credibility_label='INDUSTRY',
    credibility_color=RGBColor(0xE6, 0x5C, 0x00),
    source_name='ABC Fitness Wellness Watch Fall 2024 Report — Health & Fitness Association',
    source_type='Industry trade association report | Health & Fitness Association (formerly IHRSA)',
    url='https://www.healthandfitness.org/about/media-center/press-releases/abc-fitness-releases-wellness-watch-fall-2024-report-highlighting-generational-fitness-trends/',
    facts=[
        '54% of new gym memberships are now Gen Z',
        'Documents generational shifts in fitness facility membership across age groups',
        'Health & Fitness Association is the primary U.S. fitness industry trade body — data is widely cited',
        'Use for: gym membership demographic statistics',
    ]
)

add_source_block(doc,
    credibility_label='INDUSTRY',
    credibility_color=RGBColor(0xE6, 0x5C, 0x00),
    source_name='Demographic Data Reshaping the Fitness Industry — ABC Fitness',
    source_type='Industry report | ABC Fitness Solutions',
    url='https://abcfitness.com/abc-articles/demographic-data-reshaping-fitness-industry/',
    facts=[
        'Those under 25 accounted for 30.8% of all U.S. fitness facility members in 2024',
        'Up from 22.9% in 2015 — nearly 10 percentage point increase over 9 years',
        'Based on ABC Fitness\'s network data across hundreds of fitness facilities',
        'Use for: longitudinal trend data showing youth fitness participation growth',
    ]
)

add_source_block(doc,
    credibility_label='INDUSTRY',
    credibility_color=RGBColor(0xE6, 0x5C, 0x00),
    source_name='Generation Active: The 80% Your Club Can\'t Ignore — Les Mills',
    source_type='Industry research report | Les Mills International',
    url='https://www.lesmills.com/us/clubs-and-facilities/research-insights/audience-insights/generation-active-the-80-your-club-cant-ignore/',
    facts=[
        'Gen Z and Millennials represent 80% of all gym-goers globally',
        '73% of Gen Z have a gym membership vs ~40-50% of Gen X and Boomers',
        'Les Mills is a global fitness company operating in 100+ countries — data based on large member surveys',
        'Use for: generational gym participation statistics',
    ]
)

add_source_block(doc,
    credibility_label='INDUSTRY',
    credibility_color=RGBColor(0xE6, 0x5C, 0x00),
    source_name='Gen Z Fitness Pulse Report 2025 — The Gym Group',
    source_type='Primary research report | The Gym Group (UK)',
    url='https://www.thegymgroup.com/blog/gen-z-fitness-pulse-report-key-findings/',
    facts=[
        'Average monthly fitness spending among Gen Z rose 17% year-over-year',
        '44% of Gen Z rank fitness as their first or second spending priority — ahead of streaming and eating out',
        '79% of Gen Z use apps or digital devices to monitor their health',
        'Note: UK-based data — use for behavioral trends, note geography when citing spending figures',
        'Use for: Gen Z digital health tracking behavior and spending priorities',
    ]
)

# ─────────────────────────────────────────────
# SECTION 4: MARKET SIZE DATA
# ─────────────────────────────────────────────
add_section_header(doc, '✅ Market Size Data (use for financial projections)')

add_source_block(doc,
    credibility_label='MARKET DATA',
    credibility_color=RGBColor(0x6A, 0x1B, 0x9A),
    source_name='Digital Health Tracking App Market — Towards Healthcare',
    source_type='Market research report | Towards Healthcare',
    url='https://www.towardshealthcare.com/insights/digital-health-tracking-app-market-sizing',
    facts=[
        'Digital health tracking app market: $16.11B in 2024, grew to $18.68B in 2025',
        'Projected to reach $67.97B by 2034 — 15.94% CAGR',
        'Growth driven by wearable adoption, smartphone penetration, personal health management awareness',
        'Use for: addressable market size in your Market section',
    ]
)

add_source_block(doc,
    credibility_label='MARKET DATA',
    credibility_color=RGBColor(0x6A, 0x1B, 0x9A),
    source_name='Wellness App Market Projection 2026–2035 — Business Research Insights',
    source_type='Market research report | Business Research Insights',
    url='https://www.businessresearchinsights.com/market-reports/wellness-app-market-117356',
    facts=[
        'Global wellness app market: $3.74B in 2025, projected $15.85B by 2034 — 17.7% CAGR',
        'AI-powered wellness market specifically projected to reach $3.3B by 2027 at 36.2% CAGR',
        'North America leads with 44.5% market share',
        'Use for: wellness app market size and AI health growth rate',
    ]
)

# ─────────────────────────────────────────────
# SECTION 5: SOURCES TO REPLACE / NOT USE
# ─────────────────────────────────────────────
add_section_header(doc, '❌ Sources to Avoid / Replace', color=(0xC6, 0x28, 0x28))

p = doc.add_paragraph()
p.add_run('The following sources appeared in initial research but are not credible enough for academic submission. Replace with UW Library alternatives noted below.').font.size = Pt(10)
doc.add_paragraph('')

weak_sources = [
    ('muscleandbrawn.com', 'Fitness blog that aggregates statistics from other sources — not an original source', 'Search IHRSA/Health & Fitness Association reports via UW Library'),
    ('scrumball.com', 'Marketing company blog — no research methodology or credentials', 'Use the Tandfonline scoping review instead (already listed above)'),
    ('glofox.com', 'Gym software company blog — content marketing, not independent research', 'Use Health & Fitness Association data instead'),
]

for source, reason, replacement in weak_sources:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f'{source}  ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    r2 = p.add_run(f'— {reason}')
    r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.4)
    r3 = p2.add_run(f'→ Replace with: {replacement}')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x1e, 0x88, 0x55)
    r3.italic = True
    doc.add_paragraph('')

# ─────────────────────────────────────────────
# SECTION 6: UW LIBRARY SEARCH GUIDE
# ─────────────────────────────────────────────
add_section_header(doc, '📚 UW Library Search Guide')

searches = [
    ('PubMed / Web of Science', [
        '"n-of-1 trial" consumer health — academic backing for personal health experiments',
        '"health information seeking behavior" social media — scholarly version of influencer trend',
        '"mHealth app" engagement adherence — shows gap in existing apps (low retention)',
        '"personalized medicine" consumer adoption — market shift toward personalization',
    ]),
    ('Business Source Complete (EBSCO)', [
        '"digital health market" consumer behavior millennials',
        '"wellness industry" growth trends 2024 2025',
        'IHRSA Health Club Consumer Report — most cited annual fitness industry report',
    ]),
    ('Statista (UW likely has full access)', [
        'Search "fitness app market size" — get charts with methodology you can screenshot/cite',
        'Search "Gen Z health spending" — consumer expenditure data with sources',
        'Search "wellness app downloads" — app store data by demographic',
    ]),
]

for db, queries in searches:
    h = doc.add_paragraph()
    h.add_run(db).font.bold = True
    for q in queries:
        qp = doc.add_paragraph(style='List Bullet')
        qp.paragraph_format.left_indent = Inches(0.4)
        qr = qp.add_run(q)
        qr.font.size = Pt(10)
    doc.add_paragraph('')

output_path = '/sessions/festive-eager-euler/mnt/Healthapp/NVDP_Sources_Reference.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
